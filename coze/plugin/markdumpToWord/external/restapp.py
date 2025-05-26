from flask import Flask, request, jsonify
import markdown
from docx import Document
import tempfile
from datetime import datetime
import tos
import os

app = Flask(__name__)

@app.route('/convert', methods=['POST'])
def convert_markdown_to_docx():
    # 接收参数
    data = request.get_json()
    md_content = data.get('md')
    ak = os.getenv('TOS_ACCESS_KEY')
    sk = os.getenv('TOS_SECRET_KEY')


    # 参数校验
    if not all([md_content, ak, sk]):
        return jsonify({"error": "Missing required parameters (md, ak, sk)"}), 400

    try:
        # 创建Word文档
        doc = Document()
        
        # 解析Markdown内容
        for line in md_content.split("\n"):
            line = line.strip()
            if not line:
                continue  # 跳过空行
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="ListBullet")
            elif line.startswith("1. "):
                doc.add_paragraph(line[3:], style="ListNumber")
            elif "**" in line:
                p = doc.add_paragraph()
                run = p.add_run(line.replace("**", ""))
                run.bold = True
            elif "*" in line and not line.startswith("*"):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("*", ""))
                run.italic = True
            else:
                doc.add_paragraph(line)

        # 生成唯一文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = f'report_{timestamp}.docx'
        
        # TOS配置
        endpoint = "tos-cn-beijing.volces.com"
        region = "cn-beijing"
        bucket_name = "pub-kylin"
        base_path = "citic"
        object_key = os.path.join(base_path, output_file)
        final_url = f"https://{bucket_name}.{endpoint}/{object_key}"

        # 保存并上传文件
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as temp_file:
            doc.save(temp_file.name)
            tos_client = tos.TosClientV2(ak, sk, endpoint, region)
            try:
                tos_client.put_object_from_file(
                    bucket=bucket_name,
                    key=object_key,
                    file_path=temp_file.name
                )
            except Exception as e:
                return jsonify({"error": f"TOS upload failed: {str(e)}"}), 500

        return jsonify({"doc_url": final_url}), 200

    except Exception as e:
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
