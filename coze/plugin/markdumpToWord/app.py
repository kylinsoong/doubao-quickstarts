import markdown
import os
from docx import Document
import tempfile
from datetime import datetime
import tos

def markdown_to_docx(md_file, ak, sk):
    with open(md_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    print(type(md_content))
    html_content = markdown.markdown(md_content)

    doc = Document()
    
    for line in md_content.split("\n"):
        if line.startswith("# "):  # 标题1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  # 标题2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):  # 标题3
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):  # 无序列表
            doc.add_paragraph(line[2:], style="ListBullet")
        elif line.startswith("1. "):  # 有序列表
            doc.add_paragraph(line[3:], style="ListNumber")
        elif "**" in line:  # 加粗文本
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
        elif "*" in line:  # 斜体文本
            p = doc.add_paragraph()
            run = p.add_run(line.replace("*", ""))
            run.italic = True
        else:
            doc.add_paragraph(line)  # 普通段落

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f'report_{timestamp}.docx'
    endpoint = "tos-cn-beijing.volces.com"
    region = "cn-beijing"
    bucket_name = "pub-kylin"
    tos_client = tos.TosClientV2(ak, sk, endpoint, region)
    base_path = "citic"
    object_key = os.path.join(base_path, output_file)
    final_path = "https://pub-kylin.tos-cn-beijing.volces.com/" + object_key

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_file:
        temp_filename = temp_file.name
        doc.save(temp_filename)
        try:
            result = tos_client.put_object_from_file(bucket_name, object_key, temp_filename)
            print(result)
        except Exception as e:
            print(f"Error saving file: {e}")
   
    print(f"Markdown 文件已转换为 Word: {final_path}")

ak = os.getenv('TOS_ACCESS_KEY')
sk = os.getenv('TOS_SECRET_KEY')

markdown_to_docx("sample.md", ak, sk)
