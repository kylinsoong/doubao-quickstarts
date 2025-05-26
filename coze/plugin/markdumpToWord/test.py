import markdown
import os
from docx import Document
import tempfile
from datetime import datetime

def markdown_to_docx(md_file):
    try:
        with open(md_file, "r", encoding="utf-8") as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"Error: The file {md_file} was not found.")
        return
    except Exception as e:
        print(f"An error occurred while reading the file: {e}")
        return

    html_content = markdown.markdown(md_content)

    doc = Document()

    for line in md_content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="ListBullet")
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:], style="ListNumber")
        elif "**" in line:
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
        elif "*" in line:
            p = doc.add_paragraph()
            run = p.add_run(line.replace("*", ""))
            run.italic = True
        else:
            doc.add_paragraph(line)

    # 保存 Word 文档
    output_file = os.path.splitext(md_file)[0] + ".docx"
    doc.save(output_file)
    print(f"Converted {md_file} to {output_file} successfully.")

markdown_to_docx("example.md")
