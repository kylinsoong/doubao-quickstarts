#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
培训材料Word文档生成脚本
功能：将结构化的JSON内容转换为格式化的Word文档
"""

import sys
import json
import argparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    """
    设置中文字体和格式

    Args:
        run: 文本运行对象
        font_name: 字体名称
        font_size: 字号（磅）
        bold: 是否加粗
        color: 颜色RGB值，如RGBColor(255,0,0)
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def create_training_material(topic, audience, content_json, output_path):
    """
    创建培训材料Word文档

    Args:
        topic: 培训主题
        audience: 听众描述
        content_json: 结构化内容（JSON字符串）
        output_path: 输出文件路径

    Returns:
        bool: 成功返回True，失败返回False

    Raises:
        ValueError: JSON解析失败或必要字段缺失
    """
    try:
        # 解析JSON内容
        content = json.loads(content_json)

        # 验证必要字段
        if 'chapters' not in content:
            raise ValueError("JSON内容缺少必需字段: chapters")

        # 创建Word文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        font.size = Pt(12)

        # 添加封面信息
        if 'title' in content:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(content['title'])
            set_chinese_font(title_run, font_name='黑体', font_size=24, bold=True)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if 'subtitle' in content:
            subtitle_para = doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(content['subtitle'])
            set_chinese_font(subtitle_run, font_name='宋体', font_size=16)
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加培训主题和听众信息
        doc.add_paragraph()
        topic_para = doc.add_paragraph()
        topic_para.add_run(f"培训主题：{topic}")
        set_chinese_font(topic_para.runs[0], font_name='宋体', font_size=12)

        audience_para = doc.add_paragraph()
        audience_para.add_run(f"目标听众：{audience}")
        set_chinese_font(audience_para.runs[0], font_name='宋体', font_size=12)

        if 'author' in content:
            author_para = doc.add_paragraph()
            author_para.add_run(f"制作人：{content['author']}")
            set_chinese_font(author_para.runs[0], font_name='宋体', font_size=12)

        # 添加分页符
        doc.add_page_break()

        # 生成章节内容
        for chapter in content['chapters']:
            if 'chapter_title' not in chapter:
                continue

            # 添加章节标题
            chapter_heading = doc.add_heading(chapter['chapter_title'], level=1)
            set_chinese_font(chapter_heading.runs[0], font_name='黑体', font_size=18, bold=True)

            # 处理章节内容
            if 'sections' in chapter:
                for section in chapter['sections']:
                    # 添加小节标题
                    if 'section_title' in section:
                        section_heading = doc.add_heading(section['section_title'], level=2)
                        set_chinese_font(section_heading.runs[0], font_name='黑体', font_size=16, bold=True)

                    # 添加段落内容
                    if 'content' in section:
                        # 处理多段落
                        paragraphs = section['content'].split('\n\n')
                        for para_text in paragraphs:
                            if para_text.strip():
                                content_para = doc.add_paragraph()
                                content_run = content_para.add_run(para_text.strip())
                                set_chinese_font(content_run, font_name='宋体', font_size=12)

                    # 添加要点列表
                    if 'points' in section and isinstance(section['points'], list):
                        for point in section['points']:
                            list_para = doc.add_paragraph(point, style='List Bullet')
                            set_chinese_font(list_para.runs[0], font_name='宋体', font_size=12)

                    # 添加空行分隔
                    if 'content' in section or 'points' in section:
                        doc.add_paragraph()

        # 保存文档
        doc.save(output_path)
        print(f"培训材料已成功生成: {output_path}")
        return True

    except json.JSONDecodeError as e:
        print(f"错误: JSON解析失败 - {str(e)}")
        return False
    except Exception as e:
        print(f"错误: 文档生成失败 - {str(e)}")
        return False


def main():
    """命令行入口函数"""
    parser = argparse.ArgumentParser(description='培训材料Word文档生成器')
    parser.add_argument('--topic', required=True, help='培训主题')
    parser.add_argument('--audience', required=True, help='听众描述')
    parser.add_argument('--content', required=True, help='结构化内容（JSON字符串）')
    parser.add_argument('--output', required=True, help='输出文件路径')

    args = parser.parse_args()

    success = create_training_material(
        topic=args.topic,
        audience=args.audience,
        content_json=args.content,
        output_path=args.output
    )

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()

