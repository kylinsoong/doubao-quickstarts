import markdown
from bs4 import BeautifulSoup
from docx import Document
import os

def markdown_to_docx(markdown_text, output_path):
    # Convert Markdown to HTML
    html = markdown.markdown(markdown_text)

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')

    # Create a new Word Document
    doc = Document()

    for element in soup.contents:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'blockquote':
            doc.add_paragraph(element.text, style='Intense Quote')
        else:
            doc.add_paragraph(element.text)

    # Save the document
    doc.save(output_path)
    print(f"Word file saved to: {output_path}")

# Example usage
markdown_text = """
# DeepSeek部署与使用
企业或组织可通过以下三种方式快速部署并使用DeepSeek：云端DeepSeek SaaS服务、云端租户内私有DeepSeek方案，以及DeepSeek一体机数据中心内的私有部署。

1. **云端DeepSeek SaaS服务**：适用于希望通过API快速集成DeepSeek模型的企业。DeepSeek已在国内外多家公有云平台上线，而使用火山引擎DeepSeek服务具备以下优势：
    - **更强的承载能力**：火山引擎将初始TPM（每分钟Token数）提升至500万，远超阿里云的120万及百度智能云、硅基流动的1万。同时，火山引擎还率先推出50亿初始离线TPD（每日Token数）配额，满足企业大规模离线任务需求，降低应用门槛，支持更高并发。
    - **更快的推理速度**：火山引擎通过全栈自研推理引擎，将TPOT（吐字间隔）降至30ms左右，目标稳定在15 - 30ms区间。DeepSeek - R1模型平均延迟33.0ms，明显优于其他厂商，并具备完善的监控与告警机制，确保高流量场景下的稳定性。
    - **更优的联网搜索能力**：火山引擎DeepSeek具备联网搜索功能，避免大模型“将故事当新闻”的问题，确保回答的时效性和准确度。其支持自定义内容源及引用条数、联网意图识别和内容改写等高级配置，满足企业的个性化需求。
2. **云端租户内私有DeepSeek方案**：适用于对数据安全、模型定制化以及长期稳定使用有较高要求的企业。火山引擎提供以下两种私有化部署方式：
    - **云基础架构上直接部署**：依托GPU云服务器ECS与容器服务VKE，采用多机分布式架构部署满血版DeepSeek R1/V3。结合RDMA网络、SGLang推理引擎及LWS（Leader - Worker Set）编排机制，实现模型的高效推理。通过云上托管Prometheus监控服务和VKE的事件驱动自动扩容功能，保障推理服务的稳定性和可观测性，满足企业在高并发场景下的动态扩缩容需求。
    - **机器学习平台自动部署**：火山引擎机器学习平台为模型开发训练、推理应用提供一站式服务平台，可降低DeepSeek部署的复杂性，可通过点击几次按钮完成DeepSeek满血版部署。具体资源准备需准备云基础、镜像仓库等相关产品资源，并创建机器学习平台队列；在模型部署阶段，先创建在线服务并部署模型实例，选择VPC网络环境、API网关等进行一键部署，之后在在线服务详情页获取调用地址进行服务调用测试。
3. **DeepSeek一体机数据中心内的私有部署方案**：针对安全监管要求高、数据不能出数据中心的企业。火山引擎的AI一体机方案在功能架构、HiAgent智能体平台和模型安全性三个维度均展现出行业领先优势：
    - **全栈式功能架构实现高效闭环**：火山引擎一体机集成AI云原生基座、DeepSeek全尺寸模型、豆包大模型，以及轻量模型训练平台、企业AI应用创新平台等，涵盖从模型部署到AI应用开发的全流程能力。在大模型部署阶段，AI云原生基座实现小时级部署和资源管理；模型推理阶段，借助推理引擎优化等提升体验；模型后训练阶段，支持数据集管理和模型微调等；AI应用开发阶段，提供100 +行业插件，结合安全防火墙保障合规。
    - **HiAgent智能体平台重构开发范式**：HiAgent为开发者带来从传统代码驱动到智能体赋能的范式升级，加速大模型场景的快速落地。Hiagent零代码、低代码构建体系，内置提示词工程引擎与可视化工作流编排工具，业务人员通过拖拽组件即可完成智能体搭建；Hiagent集成行业插件生态，开放100 +行业插件（如金融合规审核、生产排程优化），支持企业快速使用行业经验，提升生产效率。
    - **大模型安全应用防火墙**：随着DeepSeek模型的大量部署使用，一些常见的隐患，例如有害信息生成（如价值观偏差内容）、敏感数据泄露（通过模型推理暴露隐私信息）以及提示词注入攻击（操控模型输出恶意结果）等成为CIO重点关注项，火山引擎一体机方案包括大模型安全防火墙组件，能够实时监控流量，获取用户输入数据并进行检测与反馈，识别并阻断有害内容，并根据识别结果提供判断反馈。 
"""

# Save to Word
output_file = "sample.docx"
markdown_to_docx(markdown_text, output_file)

